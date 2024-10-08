import datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.oxml import qn
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls
from docx.shared import Mm, Inches, Pt

from docx.oxml.shared import OxmlElement, qn
from docxtpl import DocxTemplate, InlineImage, RichText
from docx.oxml import parse_xml

from config.bot_settings import settings, logger, BASE_DIR


def format_new_doc(data):
    doc = Document()
    # Изменяем поля документа
    section = doc.sections[0]

    # Устанавливаем поля (в дюймах)
    doc.sections[0].page_width = Mm(210)
    section.top_margin = Mm(5)    # Верхнее поле
    section.bottom_margin = Mm(5)  # Нижнее поле
    section.left_margin = Mm(10)    # Левое поле
    section.right_margin = Mm(5)   # Правое поле
    # doc = doc('шаблон расчета.docx')
    img_path = BASE_DIR / 'header.png'
    doc.add_picture(img_path.as_posix(), width=Mm(190))
    styles = doc.styles
    for style in styles:
        if style.type == style.type.TABLE:
            print(style.name)

    # for paragraph in doc.paragraphs:
    #     if '{date}' in paragraph.text:
    #         paragraph.text = paragraph.text.replace("{date}", "new_value")

    p = doc.add_paragraph().add_run(text=f'Дата: {datetime.datetime.now(tz=settings.tz).date()}')
    p.bold = True
    p = doc.add_paragraph().add_run(text=f'Телефон клиента: {data["1"]}')
    p.bold = True
    # doc.add_paragraph()
    p = doc.add_paragraph().add_run(text=f'Марка автомобиля: {data["2"]}')
    p.bold = True
    p.font.size = Pt(20)

    # Коврики 3 - 9 -----------------------------------------------------------------------------
    header = ('Комплекта-ция',	'Характерис-тики', 'Цвет',	'',	'Стои-мость',	'Кол-во',	'Итого')
    # for style in styles:
    #     if style.type == style.type.TABLE and style.name not in ['Normal Table']:
    # doc.add_paragraph().add_run(style.name)
    total_price = 0
    table_price = 0
    table = doc.add_table(rows=len(data['step1']) + 1, cols=7)
    # table.style = style.name
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    # Заголовок таблицы
    for i, cell in enumerate(hdr_cells):
        hdr_cells[i].text = header[i]
        cell_font = cell.paragraphs[0].runs[0]
        cell_font.bold = True
        # cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="A9A9A9"/>'.format(nsdecls('w'))))

    # Заполнение строк таблицы
    logger.debug(f'steps1: {data["step1"]}')
    table.style = 'Table Grid'
    for row_num, step1 in enumerate(data['step1'], 1):
        logger.debug(f'Ряд {row_num}, {step1}')
        price = step1[4]
        count = step1[3]  # Введите количество комплектов
        total = round(float(price) * float(count))
        table_price += total
        logger.debug(f'price: {price}, count: {count}, total: {total}')
        row = [step1[0], step1[1], step1[2], step1[5], price, count, total]
        logger.debug(f'row: {row}')
        cells = table.rows[row_num].cells
        for i, cell in enumerate(cells):
            cell.text = str(row[i])
    total_price += table_price
    table.autofit = True
    doc.add_paragraph()
    logger.debug(f'table 1: {table_price}. Total: {total_price}')

    # Детали конфигурации и опции ----------------------------------------------------------------------------
    header = ('Детали конфигурации и опции', '', '', '', '', '', '')
    table_price = 0
    table = doc.add_table(rows=4, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    cell_to_merge = table.cell(0, 0)
    for i in range(1, 7):
        cell_to_merge.merge(table.cell(0, i))
    # Заголовок таблицы
    for i, cell in enumerate(hdr_cells):
        hdr_cells[i].text = header[i]
        cell_font = cell.paragraphs[0].runs[0]
        cell_font.bold = True

    table.style = 'Table Grid'
    # Заполнение строк таблицы
    count = data.get('11', 0)
    price = data.get('12', 0)
    total = round(float(price) * float(count))
    table_price += total
    comment = data.get('13', '')
    comment = '' if comment == 'Нет' else comment
    row = ['Площадка левой ноги', data.get('10'), '', comment, price, count, total]
    cells = table.rows[1].cells
    for i, cell in enumerate(cells):
        cell.text = str(row[i])

    count = data.get('15', 0)
    price = data.get('16', 0)
    total = round(float(price) * float(count))
    table_price += total
    comment = data.get('17', '')
    comment = '' if comment == 'Нет' else comment
    row = ['Перемычка 2го ряда', data.get('14'), '', comment, price, count, total]
    cells = table.rows[2].cells
    for i, cell in enumerate(cells):
        cell.text = str(row[i])

    count = data.get('19', 0)
    price = data.get('20', 0)
    total = round(float(price) * float(count))
    table_price += total
    comment = data.get('21', '')
    comment = '' if comment == 'Нет' else comment
    row = ['Подпятник', data.get('18'), '', comment, price, count, total]
    cells = table.rows[3].cells
    for i, cell in enumerate(cells):
        cell.text = str(row[i])
    table.autofit = True

    total_price += table_price
    logger.debug(f'table 2: {table_price}. Total: {total_price}')

    doc.add_paragraph()


    # Окантовка ----------------------------------------------------------------------
    header = ('Окантовка', '', '', '', '', '', '')
    table_price = 0
    table = doc.add_table(rows=3, cols=7)
    # table.style = style.name
    table.style = 'Table Grid'
    cell_to_merge = table.cell(0, 0)
    for i in range(1, 7):
        cell_to_merge.merge(table.cell(0, i))
    # Заголовок таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Окантовка'
    cell_font = hdr_cells[0].paragraphs[0].runs[0]
    cell_font.bold = True

    # Заполнение строк таблицы
    count = data.get('30')
    price = data.get('31')
    total = round(float(price) * float(count))
    table_price += total
    comment = data.get('32', '')
    comment = '' if comment == 'Нет' else comment
    if data.get('24') == 'Одинарная': #  Выберите тип прострочки окантовки?
        cant_color = data.get('25')
    else:
        cant_color = f"{data.get('26')}, {data.get('26')}"
    row = ['Тип окантовки', data.get('22'), data.get('23'), comment, price, count, total]
    cells = table.rows[1].cells
    for i, cell in enumerate(cells):
        cell.text = str(row[i])
    row = ['Строчка', data.get('24'), cant_color, '', '', '', '']
    cells = table.rows[2].cells
    for i, cell in enumerate(cells):
        cell.text = str(row[i])
    total_price += table_price
    table.autofit = True
    logger.debug(f'table 3: {table_price}. Total: {total_price}')
    doc.add_paragraph()

    # Вышивки и другие опции ----------------------------------------------------------------------
    header = ('Вышивки и другие опции', '', '', '', '')
    table_price = 0
    table = doc.add_table(rows=len(data['step2']) + 2, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    # Заголовок таблицы
    cell_to_merge = table.cell(0, 0)
    for i in range(1, 5):
        cell_to_merge.merge(table.cell(0, i))
    for i, cell in enumerate(hdr_cells):
        hdr_cells[i].text = header[i]
        cell_font = cell.paragraphs[0].runs[0]
        cell_font.bold = True
    # table = doc.add_table(rows=len(data['step2']) + 1, cols=5)
    table.style = 'Table Grid'

    count = data.get('34', 0)
    price = data.get('35', 0)
    total = round(float(price) * float(count))
    table_price += total
    comment = data.get('41', '')
    comment = '' if comment == 'Нет' else comment
    row = ['Вышивка логотипа', comment, price, count, total]
    cells = table.rows[1].cells
    for i, cell in enumerate(cells):
        cell.text = str(row[i])
    
    # Опции step2
    logger.debug(f'steps2: {data["step2"]}')
    for row_num, step2 in enumerate(data['step2'], 1):
        logger.debug(f'Ряд {row_num}, {step2}')
        price = step2[2]
        count = step2[1]  # Введите количество комплектов
        total = round(float(price) * float(count))
        table_price += total
        logger.debug(f'price: {price}, count: {count}, total: {total}')
        row = [step2[0], step2[1], price, count, total]
        logger.debug(f'row: {row}')
        cells = table.rows[row_num].cells
        for i, cell in enumerate(cells):
            cell.text = str(row[i])
    total_price += table_price
    logger.debug(f'table 4: {table_price}. Total: {total_price}')
    doc.add_paragraph()

    # Фото
    table = doc.add_table(rows=3, cols=5)
    hdr_cells = table.rows[0].cells
    # Заголовок таблицы
    cell_to_merge = table.cell(0, 0)
    for i in range(1, 5):
        cell_to_merge.merge(table.cell(0, i))
    for i, cell in enumerate(hdr_cells):
        hdr_cells[i].text = header[i]
        cell_font = cell.paragraphs[0].runs[0]
        cell_font.bold = True
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Фото'
    # table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    step3 = data['step3']
    order_id = data['order_id']
    for i, photo_pair in enumerate(step3):
        table.rows[1].cells[i].paragraphs[0].add_run().add_picture(f"photo_{order_id}_{i}.jpg", width=Mm(35))
        table.rows[2].cells[i].text = photo_pair[1]
    doc.add_paragraph()

    # Итого
    table = doc.add_table(rows=1, cols=3)
    table_price = 0
    # table.style = 'Table Grid'
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # page_width = int(doc.sections[0].page_width.mm) - int(doc.sections[0].left_margin.mm) - int(doc.sections[0].right_margin.mm) + 3
    page_width = 200

    row_cells = table.rows[0].cells
    for i in range(3):
        cell = row_cells[i]
        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w'))))

    cell_to_merge = table.cell(0, 0)
    cell_to_merge.merge(table.cell(0, 1))
    cell_to_merge.text = 'ИТОГО'
    p = cell_to_merge.paragraphs[0].runs[0]
    p.font.size = Pt(16)
    p.bold = True
    cell = table.cell(0, 2)
    cell.text = f'{total_price}'

    def set_cell_border(cell, **kwargs):
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
            bottom={"sz": 12, "color": "00FF00", "val": "single"},
            left={"sz": 24, "val": "dashed", "shadow": "true"},
            right={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    discount = float(data.get("47", 0))
    if discount:
        table.add_row()
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        row_cells = table.rows[1].cells
        for i in range(1, 3):
            cell = row_cells[i]
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="00CA00"/>'.format(nsdecls('w'))))
        cell = table.cell(1, 1)
        cell.text = f'ИТОГО с учетом %'
        cell = table.cell(1, 2)
        cell.text = f'{round(total_price - total_price * discount / 100)}'
    else:
        row = table.add_row()
        for i in range(3):
            set_cell_border(table.cell(0, i),
                            bottom={"sz": 6, "color": "#000000", "val": "single"},
                            top={"sz": 6, "color": "#000000", "val": "single"},
                            left={"sz": 6, "color": "#000000", "val": "single"},
                            right={"sz": 6, "color": "#000000", "val": "single"},
                            )
        for row in table.rows:
            for cell in row.cells:
                pass

    # Устанавливаем ширину для последнего столбца

    last_cell_width = 30
    cell_width = int((page_width - last_cell_width) / 2)
    x = 30
    widths = (Mm(cell_width + x), Mm(cell_width - x -2), Mm(last_cell_width - 3))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # tables = doc.tables[:-2]
    # for table in tables:
    #     rows = table.rows
    #     for i, row in enumerate(rows):
    #         cells = row.cells
    #         for n in range(-1, -4, -1):
    #             cell = cells[n]
    #             cell.width = Mm(20)
    #     table.autofit = True

    doc.save('demo.docx')
    return doc


if __name__ == '__main__':
    data = {'count': 48, 'step1': [['Задние коврики', 'Prime 7мм', 'Синий', '0.6', '10000', '']], 'step2': [], 'step3': [['AgACAgIAAxkBAAJu42bVGKcj6c9mbmElcW-ZmzwzHTlpAAKw2TEbDBGxStS8P2YkvPaXAQADAgADeQADNQQ', 'Коммент1'], ['AgACAgIAAxkBAAJu6WbVGLoujKT3vukhzE0sH33ke_hZAAKx2TEbDBGxSlIBEeSyZRzuAQADAgADeQADNQQ', 'Коммент2']], 'order_id': 1, '1': '123', '2': 'Dongfeng', '3': 'Задние коврики', '4': 'Prime 7мм', '5': 'Синий', '6': '0.6', '7': '10000', '8': 'Нет', '9': 'Нет', '10': 'Без язычка', '14': 'Без перемычки', '18': 'Без подпятника', '22': 'Экокожи с подгибкой', '23': 'Темно-серый', '24': 'Одинарная', '25': 'Темно-серый', '28': 'Нет', '30': '0.6', '31': '123', '32': 'Нет', '33': 'Нет', '37': 'Нет', '43': 'Да', '44': 'AgACAgIAAxkBAAJu6WbVGLoujKT3vukhzE0sH33ke_hZAAKx2TEbDBGxSlIBEeSyZRzuAQADAgADeQADNQQ', '45': 'Коммент2', '46': 'Нет', '47': '5'}
    # data = {'count': 48, 'step1': [['Водительский коврик', 'Brilliance 12мм', 'Темно-серый', '0.3', '10000', '']], 'step2': [],
    #  'step3': [],  '1': '124', '2': 'Smart', '3': 'Водительский коврик', '4': 'Brilliance 12мм', '5': 'Темно-серый', '6': '0.3', '7': '10000', '8': 'Нет', '9': 'Нет', '10': 'Без язычка', '14': 'Без перемычки', '18': 'Без подпятника', '22': 'Текстильная', '23': 'Серый', '24': 'Одинарная', '25': 'Темно-серый', '28': 'Нет', '30': '0.3', '31': '2134', '32': 'Нет', '33': 'Нет', '37': 'Нет', '43': 'Нет', '47': '5'}
    format_new_doc(data)
    # from docx2pdf import convert

    # convert("demo.docx")

