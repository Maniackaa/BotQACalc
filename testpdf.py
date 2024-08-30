import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Mm

from config.bot_settings import settings, logger
from database.db import get_price, get_count

doc = Document()

# Изменяем поля документа
section = doc.sections[0]

# Устанавливаем поля (в дюймах)
section.top_margin = Mm(5)    # Верхнее поле
section.bottom_margin = Mm(5)  # Нижнее поле
section.left_margin = Mm(5)    # Левое поле
section.right_margin = Mm(5)   # Правое поле
# doc = doc('шаблон расчета.docx')
doc.add_picture('header.png', width=Mm(200))
styles = doc.styles
data = {'count': 8, 'step1': [['Коврик в багажник', 'Luxury 20мм', 'Серый', 'коммент', 'Нет']], 'q_content_type':'text', '1': '1234124124', '2': 'JAC', '3': 'Коврик в багажник', '4': 'Luxury 20мм', '5': 'Серый', '6': 'коммент', '7': 'Нет'}
for paragraph in doc.paragraphs:
    if '{date}' in paragraph.text:
        print(paragraph.text)
        paragraph.text = paragraph.text.replace("{date}", "new_value")

p = doc.add_paragraph().add_run(text=f'Дата: {datetime.datetime.now(tz=settings.tz).date()}')
p.bold = True
p = doc.add_paragraph().add_run(text=f'Телефон клиента: {data["1"]}')
p.bold = True
doc.add_paragraph()
p = doc.add_paragraph().add_run(text=f'Марка автомобиля: {data["2"]}')
p.bold = True


header = ('Комплектация',	'Характеристики', 'Цвет',	'',	'Стоимость',	'Кол-во',	'Итого')

for style in styles:
    if style.type == style.type.TABLE:
        doc.add_paragraph('-------------------')
        doc.add_paragraph(style.name)
        table = doc.add_table(rows=6, cols=7)
        table.style = style.name
        hdr_cells = table.rows[0].cells
        # Заголовок таблицы
        for i, cell in enumerate(hdr_cells):
            hdr_cells[i].text = header[i]
            cell_font = cell.paragraphs[0].runs[0]
            cell_font.bold = True
            # cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="A9A9A9"/>'.format(nsdecls('w'))))

        # Заполнение строк таблицы
        for row_num, step1 in enumerate(data['step1'], 1):
            logger.debug(f'Ряд {row_num}, {step1}')
            price = get_price(step1[1])
            count = get_count(step1[0])
            total = price * count
            logger.debug(f'price: {price}, count: {count}, total: {total}')

            row = [step1[0], step1[1], step1[2], '', price, count, total]
            logger.debug(f'row: {row}')
            cells = table.rows[row_num].cells
            for i, cell in enumerate(cells):
                cell.text = str(row[i])

        p = doc.add_paragraph('-------------------')

p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True


# doc.add_picture('photo.jpg', width=Inches(6.25))
# doc.add_page_break()

doc.save('demo.docx')


# from docx2pdf import convert
#
# convert("demo.docx", "output.pdf")